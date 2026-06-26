#!/usr/bin/env python3
"""Generate a modern Word (.docx) documentation for a TestStand sequence file.

Usage:
    py generate_teststand_doc.py <data.json> <output.docx>
        [--diagram-out <png>] [--browser <exe>] [--teststand-dir <dir>]

Requirements: python-docx (required); Pillow (optional — enables the original
TestStand step icons, tinted monochrome in the document accent color; without
Pillow or without a TestStand installation the step listing is rendered
without icons). A Chromium browser (Edge/Chrome) must be installed for
rendering the dependency diagram (SVG -> PNG headless). The TestStand
installation is auto-discovered under "Program Files*\\National
Instruments\\TestStand*" (newest version wins); override with
--teststand-dir or the TS_DOC_TESTSTAND_DIR environment variable.

JSON contract (UTF-8):
{
  "title":     "TFW_ExampleModule",         # document title (usually file name w/o extension)
  "language":  "de",                        # "de" or "en" — label set used in the document
  "generated": "2026-07-02",                # date shown in the header line
  "file": {
    "path":        "C:\\Seq\\File.seq",
    "description": "Short description of the sequence file.",
    "version":     "1.0.0.0"                # optional
  },
  "sequences": [
    {
      "name":        "MainSequence",
      "description": "What this sequence does.",
      "parameters": [
        { "name": "VoltageLimit", "type": "Number", "default": "24",
          "by_ref": false, "description": "Upper voltage limit." }
      ],
      "steps": {                             # optional: ordered steps per step group
        "Setup":   [],
        "Main":    [
          { "name": "While_Attempts", "type": "NI_Flow_While",
            "detail": "Locals.Attempt < 3" },                              # detail is optional
          { "name": "Measure", "type": "SequenceCall",
            "detail": "→ Measure_Pressure" },
          { "name": "End_While", "type": "NI_Flow_End" },
          { "name": "Old_Check", "type": "PassFailTest", "enabled": false } # rendered struck-through
        ],
        "Cleanup": []
      },
      "calls": [                             # SequenceCall steps found in this sequence
        { "target_sequence": "Init", "target_file": "", "count": 1 },      # "" = same file
        { "target_sequence": "Open", "target_file": "Driver.seq" },        # external file
        { "step_name": "Do_Later", "unresolved": true }                    # unlinked placeholder
      ]
    }
  ]
}
"""
import json
import os
import shutil
import subprocess
import sys
import tempfile
from xml.sax.saxutils import escape as xml_escape

try:
    from PIL import Image
except ImportError:  # without Pillow the step icons are silently omitted
    Image = None

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

try:
    from docx.oxml import OxmlElement
except ImportError:  # python-docx >= 1.2 moved it
    from docx.oxml.parser import OxmlElement

sys.setrecursionlimit(5000)

FONT = "Segoe UI"

# Restrained, modern palette — dark petrol accent, light fills, no icons.
C_INK = "333333"
C_MUTED = "6E7B84"
C_ACCENT = "1F4E5F"
C_ACCENT2 = "2E6E80"
C_RULE = "C9D6DC"
C_TBL_LINE = "D9E2E7"
C_TBL_LINE2 = "B9C6CD"
C_NODE_FILL = "F2F6F8"
C_NODE_STROKE = "9FB6C1"
C_EDGE = "5B7A8C"

LABELS = {
    "de": {
        "toc": "Inhaltsverzeichnis",
        "toc_placeholder": "Das Inhaltsverzeichnis wird beim ersten Öffnen in Word aktualisiert (F9).",
        "sequences": "Sequenzen",
        "dependencies": "Sequenzabhängigkeiten",
        "parameters": "Parameter",
        "no_parameters": "Keine Parameter.",
        "no_description": "Keine Beschreibung hinterlegt.",
        "param_name": "Name",
        "param_type": "Typ",
        "param_passing": "Übergabe",
        "param_default": "Standardwert",
        "param_desc": "Beschreibung",
        "by_ref": "By Reference",
        "by_val": "By Value",
        "calls": "Ruft auf",
        "unresolved_calls": "Nicht verlinkte Aufrufe",
        "version": "Version",
        "sequences_count": "Sequenzen",
        "dep_intro": "Das folgende Diagramm zeigt die Aufrufbeziehungen zwischen den Sequenzen dieser Datei{ext}.",
        "dep_intro_ext": " sowie die Abhängigkeiten zu externen Sequenzdateien",
        "dep_none": "Zwischen den Sequenzen dieser Datei bestehen keine Aufrufbeziehungen.",
        "legend_entry": "Einstiegssequenz",
        "legend_seq": "Sequenz",
        "legend_ext": "Externe Datei",
        "legend_missing": "Nicht gefunden",
        "legend_calls": "ruft auf",
        "page": "Seite",
        "page_of": "von",
        "not_found": "nicht gefunden",
        "recursive": "rekursiv",
        "steps": "Steps",
    },
    "en": {
        "toc": "Table of Contents",
        "toc_placeholder": "The table of contents is updated when the document is first opened in Word (F9).",
        "sequences": "Sequences",
        "dependencies": "Sequence Dependencies",
        "parameters": "Parameters",
        "no_parameters": "No parameters.",
        "no_description": "No description available.",
        "param_name": "Name",
        "param_type": "Type",
        "param_passing": "Passing",
        "param_default": "Default",
        "param_desc": "Description",
        "by_ref": "By Reference",
        "by_val": "By Value",
        "calls": "Calls",
        "unresolved_calls": "Unlinked calls",
        "version": "Version",
        "sequences_count": "sequences",
        "dep_intro": "The following diagram shows the call relationships between the sequences of this file{ext}.",
        "dep_intro_ext": " as well as the dependencies on external sequence files",
        "dep_none": "There are no call relationships between the sequences of this file.",
        "legend_entry": "Entry sequence",
        "legend_seq": "Sequence",
        "legend_ext": "External file",
        "legend_missing": "Not found",
        "legend_calls": "calls",
        "page": "Page",
        "page_of": "of",
        "not_found": "not found",
        "recursive": "recursive",
        "steps": "Steps",
    },
}

BROWSER_CANDIDATES = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
]


# --------------------------------------------------------------------------
# Step listing: flow-aware indentation
# --------------------------------------------------------------------------

FLOW_OPENERS = {"NI_Flow_If", "NI_Flow_While", "NI_Flow_DoWhile", "NI_Flow_For",
                "NI_Flow_ForEach", "NI_Flow_SweepLoop", "NI_Flow_StreamLoop",
                "NI_Flow_Select", "NI_Flow_Case"}
FLOW_MID = {"NI_Flow_ElseIf", "NI_Flow_Else"}
FLOW_END = {"NI_Flow_End"}
FLOW_ALL = FLOW_OPENERS | FLOW_MID | FLOW_END | {"NI_Flow_Break", "NI_Flow_Continue"}

TYPE_DISPLAY = {
    "NI_Flow_If": "If", "NI_Flow_ElseIf": "Else If", "NI_Flow_Else": "Else",
    "NI_Flow_End": "End", "NI_Flow_While": "While", "NI_Flow_DoWhile": "Do While",
    "NI_Flow_For": "For", "NI_Flow_ForEach": "For Each",
    "NI_Flow_SweepLoop": "Sweep Loop", "NI_Flow_StreamLoop": "Stream Loop",
    "NI_Flow_Select": "Select", "NI_Flow_Case": "Case",
    "NI_Flow_Break": "Break", "NI_Flow_Continue": "Continue",
    "NI_Wait": "Wait", "SequenceCall": "Sequence Call",
    "MessagePopup": "Message Popup", "CallExecutable": "Call Executable",
    "PassFailTest": "Pass/Fail Test", "NumericLimitTest": "Numeric Limit Test",
    "StringValueTest": "String Value Test",
    "NI_MultipleNumericLimitTest": "Multiple Numeric Limit Test",
}


def type_display(t):
    return TYPE_DISPLAY.get(t, t)


def step_rows(steps):
    """Yield (indent_depth, step) — NI_Flow_* blocks indent their content;
    Else/ElseIf sit on the opener's level, each End closes one level."""
    depth, rows = 0, []
    for st in steps:
        t = (st.get("type") or "").strip()
        if t in FLOW_END:
            depth = max(0, depth - 1)
            d = depth
        elif t in FLOW_MID:
            d = max(0, depth - 1)
        else:
            d = depth
        rows.append((d, st))
        if t in FLOW_OPENERS:
            depth += 1
    return rows


# --------------------------------------------------------------------------
# Original TestStand step icons, tinted monochrome (silhouette in one color)
# --------------------------------------------------------------------------

ICON_FILES = {
    "NI_Flow_If": r"Icons\FlowControl\NI_If.ico",
    "NI_Flow_ElseIf": r"Icons\FlowControl\NI_ElseIf.ico",
    "NI_Flow_Else": r"Icons\FlowControl\NI_Else.ico",
    "NI_Flow_End": r"Icons\FlowControl\NI_End.ico",
    "NI_Flow_While": r"Icons\FlowControl\NI_While.ico",
    "NI_Flow_DoWhile": r"Icons\FlowControl\NI_DoWhile.ico",
    "NI_Flow_For": r"Icons\FlowControl\NI_For.ico",
    "NI_Flow_ForEach": r"Icons\FlowControl\NI_ForEach.ico",
    "NI_Flow_SweepLoop": r"Icons\FlowControl\NI_Sweep.ico",
    "NI_Flow_StreamLoop": r"Icons\FlowControl\NI_StreamLoop.ico",
    "NI_Flow_Select": r"Icons\FlowControl\NI_Select.ico",
    "NI_Flow_Case": r"Icons\FlowControl\NI_Case.ico",
    "NI_Flow_Break": r"Icons\FlowControl\NI_Break.ico",
    "NI_Flow_Continue": r"Icons\FlowControl\NI_Continue.ico",
    "SequenceCall": r"Icons\SeqAdp.ico",
    "Statement": r"Icons\statement.ico",
    "NI_Wait": r"StepTypes\SyncSteps\res\wait.ico",
    "MessagePopup": r"Icons\MsgBox.ico",
    "CallExecutable": r"Icons\execstep.ico",
    "Action": r"Icons\NoneAdp_action.ico",
    "PassFailTest": r"Icons\NoneAdp_test.ico",
    "NumericLimitTest": r"Icons\NoneAdp_test.ico",
    "StringValueTest": r"Icons\NoneAdp_test.ico",
    "NI_MultipleNumericLimitTest": r"Icons\NoneAdp_test.ico",
    "Goto": r"Icons\goto.ico",
    "Label": r"Icons\label.ico",
}
DEFAULT_ICON = r"Icons\Generic.ico"


def find_teststand_components(override=None):
    """Locate <TestStand>\\Components (icons live below it). Newest install wins."""
    import glob as _glob
    for cand in (override, os.environ.get("TS_DOC_TESTSTAND_DIR")):
        if not cand:
            continue
        c = cand if os.path.basename(cand).lower() == "components" \
            else os.path.join(cand, "Components")
        if os.path.isdir(c):
            return c
    hits = []
    for root in (r"C:\Program Files", r"C:\Program Files (x86)"):
        hits += _glob.glob(os.path.join(root, "National Instruments",
                                        "TestStand*", "Components"))
    return sorted(hits)[-1] if hits else None


class IconTinter:
    """Renders a step type's original TestStand icon as a flat, single-color
    silhouette PNG (32 px) and caches it per (type, color)."""

    def __init__(self, components, cache_dir):
        self.components = components
        self.cache_dir = cache_dir
        self.cache = {}
        self.ok = bool(components) and Image is not None and bool(cache_dir)

    def png(self, step_type, color_hex):
        if not self.ok:
            return None
        key = (step_type, color_hex)
        if key in self.cache:
            return self.cache[key]
        src = os.path.join(self.components, ICON_FILES.get(step_type, DEFAULT_ICON))
        if not os.path.isfile(src):
            src = os.path.join(self.components, DEFAULT_ICON)
        path = None
        if os.path.isfile(src):
            try:
                im = Image.open(src).convert("RGBA").resize((32, 32), Image.LANCZOS)
                rgb = tuple(int(color_hex[i:i + 2], 16) for i in (0, 2, 4))
                out = Image.new("RGBA", im.size)
                px, po = im.load(), out.load()
                for y in range(im.height):
                    for x in range(im.width):
                        po[x, y] = (rgb[0], rgb[1], rgb[2], px[x, y][3])
                safe = "".join(c if c.isalnum() else "_" for c in step_type)
                path = os.path.join(self.cache_dir, f"{safe}_{color_hex}.png")
                out.save(path)
            except Exception:
                path = None
        self.cache[key] = path
        return path


# --------------------------------------------------------------------------
# Dependency graph: build, layer, order, position
# --------------------------------------------------------------------------

class Node:
    def __init__(self, nid, label, kind):
        self.id = nid
        self.label = label
        self.kind = kind          # "seq" | "ext" | "missing"
        self.is_root = False
        self.called = set()       # ext nodes: called sequence names (subtitle)
        self.subtitle = ""
        self.layer = 0
        self.vrow = 0
        self.x = self.y = 0.0
        self.w = self.h = 0.0


def build_graph(data, labels):
    file_base = os.path.basename(data["file"]["path"]).lower()
    nodes, order = {}, []

    def add(nid, label, kind):
        if nid not in nodes:
            nodes[nid] = Node(nid, label, kind)
            order.append(nid)
        return nodes[nid]

    for s in data["sequences"]:
        add("seq:" + s["name"], s["name"], "seq")

    edges = {}  # (src, dst) -> call count
    for s in data["sequences"]:
        src = "seq:" + s["name"]
        for c in s.get("calls", []):
            if c.get("unresolved"):
                continue
            tseq = (c.get("target_sequence") or "").strip()
            tfile = (c.get("target_file") or "").strip()
            if not tseq:
                continue
            if tfile and os.path.basename(tfile).lower() == file_base:
                tfile = ""  # a call back into the documented file itself
            if tfile:
                base = os.path.basename(tfile)
                node = add("ext:" + base.lower(), base, "ext")
                node.called.add(tseq)
                dst = node.id
            else:
                dst = "seq:" + tseq
                if dst not in nodes:
                    n = add(dst, tseq, "missing")
                    n.subtitle = "(" + labels["not_found"] + ")"
            edges[(src, dst)] = edges.get((src, dst), 0) + max(1, int(c.get("count", 1) or 1))

    for n in nodes.values():
        if n.kind == "ext":
            sub = ", ".join(sorted(n.called))
            n.subtitle = sub if len(sub) <= 46 else sub[:43] + "…"

    # Detect back edges (cycles / recursion) so layering stays acyclic.
    adj = {nid: [] for nid in order}
    for (a, b) in edges:
        adj[a].append(b)
    state, back = {}, set()

    def dfs(u):
        state[u] = 1
        for v in adj[u]:
            if state.get(v, 0) == 1:
                back.add((u, v))
            elif state.get(v, 0) == 0:
                dfs(v)
        state[u] = 2

    for nid in order:
        if state.get(nid, 0) == 0:
            dfs(nid)

    dag = [(a, b) for (a, b) in edges if (a, b) not in back]

    # Longest-path layering by relaxation (graphs are small).
    layer = {nid: 0 for nid in order}
    for _ in range(len(order)):
        changed = False
        for (a, b) in dag:
            if layer[b] < layer[a] + 1:
                layer[b] = layer[a] + 1
                changed = True
        if not changed:
            break
    for nid in order:
        nodes[nid].layer = layer[nid]

    indeg = {nid: 0 for nid in order}
    outdeg = {nid: 0 for nid in order}
    for (a, b) in dag:
        indeg[b] += 1
        outdeg[a] += 1
    for n in nodes.values():
        n.is_root = n.kind == "seq" and indeg[n.id] == 0 and outdeg[n.id] > 0

    return nodes, order, edges, back


def layout_graph(nodes, order, edges, back):
    NODE_H, SUB_H, VGAP, HGAP, MARGIN = 56, 74, 110, 46, 36
    LEGEND_H = 58

    layers = {}
    for nid in order:
        layers.setdefault(nodes[nid].layer, []).append(nid)
    max_layer = max(layers) if layers else 0

    preds, succs = {}, {}
    for (a, b) in edges:
        if (a, b) in back:
            continue
        succs.setdefault(a, []).append(b)
        preds.setdefault(b, []).append(a)

    # Barycenter ordering: two down/up sweeps reduce edge crossings.
    def sweep(get_neighbors, layer_range):
        for lyr in layer_range:
            row = layers.get(lyr, [])
            pos = {nid: i for l in layers.values() for i, nid in enumerate(l)}
            def key(nid):
                ns = get_neighbors.get(nid, [])
                return sum(pos.get(m, 0) for m in ns) / len(ns) if ns else pos.get(nid, 0)
            row.sort(key=key)

    for _ in range(2):
        sweep(preds, range(1, max_layer + 1))
        sweep(succs, range(max_layer - 1, -1, -1))

    def node_width(n):
        w = 30 + 9.3 * len(n.label)
        if n.subtitle:
            w = max(w, 26 + 7.2 * len(n.subtitle))
        return max(176.0, min(w, 480.0))

    for n in nodes.values():
        n.w = node_width(n)
        n.h = SUB_H if n.subtitle else NODE_H

    # Wrap wide layers into several centered sub-rows (a grid) so a large
    # fan-out (e.g. one main sequence calling 15 subsequences) does not
    # produce a single enormous, cluttered row. Columns adapt to the widest
    # layer: the more nodes, the more we wrap toward a squarish block.
    widest = max((len(r) for r in layers.values()), default=1)
    import math as _math
    max_cols = max(3, min(6, _math.ceil(_math.sqrt(widest * 1.6)))) if widest > 6 else widest
    MAX_ROW_W = 1500.0

    visual_rows = []
    for lyr in sorted(layers):
        chunk, chunk_w = [], 0.0
        for nid in layers[lyr]:
            w = nodes[nid].w
            if chunk and (len(chunk) >= max_cols or chunk_w + HGAP + w > MAX_ROW_W):
                visual_rows.append(chunk)
                chunk, chunk_w = [], 0.0
            chunk_w += (HGAP if chunk else 0) + w
            chunk.append(nid)
        if chunk:
            visual_rows.append(chunk)

    row_widths = [sum(nodes[nid].w for nid in vr) + HGAP * (len(vr) - 1)
                  for vr in visual_rows]
    content_w = max(row_widths + [560.0])
    W = content_w + 2 * MARGIN

    y = float(MARGIN)
    for ri, (vr, rw) in enumerate(zip(visual_rows, row_widths)):
        rh = max((nodes[nid].h for nid in vr), default=NODE_H)
        x = MARGIN + (content_w - rw) / 2.0
        for nid in vr:
            n = nodes[nid]
            n.x, n.y, n.vrow = x, y, ri
            x += n.w + HGAP
        y += rh + VGAP
    H = y - VGAP + MARGIN + LEGEND_H

    # Reserve side "lanes" for edges that cannot go straight down into their
    # target without crossing a bubble: recursion, upward/sideways calls, or
    # calls that skip a visual row. Those get routed around through the margins
    # (see make_svg), so the arrow lines never pass under a node.
    def _vr(nid):
        return nodes[nid].vrow
    wrapped = widest > max_cols
    need = wrapped or any(
        a == b or (a, b) in back or _vr(b) <= _vr(a) or _vr(b) - _vr(a) >= 2
        for (a, b) in edges)
    pad = 0.0
    if need:
        pad = 190.0
        for n in nodes.values():
            n.x += pad
        W = content_w + 2 * MARGIN + 2 * pad
    return W, H, MARGIN, pad


# --------------------------------------------------------------------------
# SVG diagram
# --------------------------------------------------------------------------

def make_svg(nodes, order, edges, back, W, H, margin, labels, scale=2, route_pad=0.0):
    parts = []
    parts.append(
        f'<svg width="{int(W * scale)}" height="{int(H * scale)}" viewBox="0 0 {int(W)} {int(H)}" '
        f'xmlns="http://www.w3.org/2000/svg" font-family="Segoe UI, sans-serif">'
    )
    parts.append(
        '<defs><marker id="arrow" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="7" '
        'markerHeight="7" orient="auto-start-reverse">'
        f'<path d="M 0 1 L 9 5 L 0 9 z" fill="#{C_EDGE}"/></marker></defs>'
    )
    parts.append(f'<rect x="0" y="0" width="{int(W)}" height="{int(H)}" fill="#ffffff"/>')

    # Obstacle-aware edge routing: a straight downward curve is used only when
    # it stays clear of every other bubble; otherwise the edge is routed around
    # through a side margin lane and the empty channels between rows, so an
    # arrow line never passes underneath a node.
    PADB = 8.0
    boxes = {nid: (n.x - PADB, n.y - PADB, n.w + 2 * PADB, n.h + 2 * PADB)
             for nid, n in nodes.items()}

    def clear(pts, a, b):
        for (px, py) in pts:
            for nid, (bx, by, bw, bh) in boxes.items():
                if nid == a or nid == b:
                    continue
                if bx <= px <= bx + bw and by <= py <= by + bh:
                    return False
        return True

    def bez_pts(p, n=48):
        out = []
        for i in range(n + 1):
            t = i / n
            mt = 1 - t
            out.append((
                mt**3 * p[0][0] + 3*mt*mt*t*p[1][0] + 3*mt*t*t*p[2][0] + t**3*p[3][0],
                mt**3 * p[0][1] + 3*mt*mt*t*p[1][1] + 3*mt*t*t*p[2][1] + t**3*p[3][1],
            ))
        return out

    def rounded(points, r=14):
        d = [f"M {points[0][0]:.0f} {points[0][1]:.0f}"]
        for i in range(1, len(points) - 1):
            (x0, y0), (x1, y1), (x2, y2) = points[i - 1], points[i], points[i + 1]

            def pull(xb, yb):
                dx, dy = xb - x1, yb - y1
                L = (dx * dx + dy * dy) ** 0.5 or 1
                rr = min(r, L / 2)
                return x1 + dx / L * rr, y1 + dy / L * rr
            ex, ey = pull(x0, y0)
            sx, sy = pull(x2, y2)
            d.append(f"L {ex:.0f} {ey:.0f}")
            d.append(f"Q {x1:.0f} {y1:.0f} {sx:.0f} {sy:.0f}")
        d.append(f"L {points[-1][0]:.0f} {points[-1][1]:.0f}")
        return " ".join(d)

    left_x = margin + route_pad * 0.55
    right_x = W - margin - route_pad * 0.55
    lane = {"L": 0, "R": 0}

    # Edges first (below the nodes).
    for (a, b), count in sorted(edges.items()):
        na, nb = nodes[a], nodes[b]
        is_back = (a, b) in back
        dash = ' stroke-dasharray="7 5"' if is_back else ""
        label_xy = None

        if a == b:
            # Recursion: compact loop at the node's right edge.
            x, cy = na.x + na.w, na.y + na.h / 2
            path = (f"M {x:.0f} {cy - 11:.0f} C {x + 46:.0f} {cy - 14:.0f}, "
                    f"{x + 46:.0f} {cy + 14:.0f}, {x + 6:.0f} {cy + 11:.0f}")
            label_xy = (x + 52, cy + 4)
        else:
            path = None
            if not is_back and nb.vrow > na.vrow:
                sx, sy = na.x + na.w / 2, na.y + na.h
                tx, ty = nb.x + nb.w / 2, nb.y
                dy = ty - sy
                ctrl = [(sx, sy), (sx, sy + dy * 0.45), (tx, ty - dy * 0.45), (tx, ty)]
                if clear(bez_pts(ctrl), a, b):
                    path = (f"M {sx:.0f} {sy:.0f} C {sx:.0f} {sy + dy*0.45:.0f}, "
                            f"{tx:.0f} {ty - dy*0.45:.0f}, {tx:.0f} {ty:.0f}")
                    label_xy = ((sx + tx) / 2 + 8, (sy + ty) / 2)
            if path is None and route_pad < 1:
                # No lane room reserved (rare) — accept a straight curve.
                sx, sy = na.x + na.w / 2, na.y + na.h
                tx, ty = nb.x + nb.w / 2, nb.y
                dy = (ty - sy) or 1
                path = (f"M {sx:.0f} {sy:.0f} C {sx:.0f} {sy + dy*0.45:.0f}, "
                        f"{tx:.0f} {ty - dy*0.45:.0f}, {tx:.0f} {ty:.0f}")
                label_xy = ((sx + tx) / 2 + 8, (sy + ty) / 2)
            elif path is None:
                # Route around via a side margin lane through empty row channels.
                downward = nb.vrow >= na.vrow
                mid = (na.x + na.w / 2 + nb.x + nb.w / 2) / 2
                side = "R" if mid >= W / 2 else "L"
                k = lane[side]
                lane[side] += 1
                base = right_x if side == "R" else left_x
                lx = base - k * 24 if side == "R" else base + k * 24
                sx = na.x + na.w / 2
                sy = na.y + na.h if downward else na.y
                ch_s = sy + 28 if downward else sy - 28
                tx = nb.x + nb.w / 2
                ty = nb.y if downward else nb.y + nb.h
                ch_t = ty - 28 if downward else ty + 28
                path = rounded([(sx, sy), (sx, ch_s), (lx, ch_s),
                                (lx, ch_t), (tx, ch_t), (tx, ty)])
                label_xy = (lx, (ch_s + ch_t) / 2)

        parts.append(
            f'<path d="{path}" fill="none" stroke="#{C_EDGE}" stroke-width="1.8"{dash} '
            'marker-end="url(#arrow)"/>'
        )
        if count > 1 and label_xy:
            parts.append(
                f'<text x="{label_xy[0]:.0f}" y="{label_xy[1]:.0f}" font-size="11.5" fill="#{C_MUTED}" '
                f'paint-order="stroke" stroke="#ffffff" stroke-width="3">×{count}</text>'
            )

    # Nodes.
    for nid in order:
        n = nodes[nid]
        if n.is_root:
            box = f'fill="#{C_ACCENT}"'
            main_fill, sub_fill, weight = "ffffff", "cfe0e6", 600
        elif n.kind == "ext":
            box = f'fill="#ffffff" stroke="#A9B7BF" stroke-width="1.5" stroke-dasharray="6 4"'
            main_fill, sub_fill, weight = "4A5A63", "7C8A92", 600
        elif n.kind == "missing":
            box = f'fill="#F7F9FA" stroke="#C2CCD1" stroke-width="1.5" stroke-dasharray="6 4"'
            main_fill, sub_fill, weight = "7C8A92", "9AA6AC", 400
        else:
            box = f'fill="#{C_NODE_FILL}" stroke="#{C_NODE_STROKE}" stroke-width="1.5"'
            main_fill, sub_fill, weight = "22333B", "7C8A92", 400
        parts.append(f'<rect x="{n.x:.0f}" y="{n.y:.0f}" width="{n.w:.0f}" height="{n.h:.0f}" rx="10" {box}/>')
        ty = n.y + (n.h / 2 + 5.5 if not n.subtitle else 30)
        parts.append(
            f'<text x="{n.x + n.w / 2:.0f}" y="{ty:.0f}" text-anchor="middle" font-size="15.5" '
            f'font-weight="{weight}" fill="#{main_fill}">{xml_escape(n.label)}</text>'
        )
        if n.subtitle:
            parts.append(
                f'<text x="{n.x + n.w / 2:.0f}" y="{n.y + 52:.0f}" text-anchor="middle" '
                f'font-size="11.5" fill="#{sub_fill}">{xml_escape(n.subtitle)}</text>'
            )

    # Minimal legend (only the kinds that actually occur).
    ly = H - 40
    lx = float(margin)
    items = []
    if any(n.is_root for n in nodes.values()):
        items.append(("swatch", f'fill="#{C_ACCENT}"', labels["legend_entry"]))
    items.append(("swatch", f'fill="#{C_NODE_FILL}" stroke="#{C_NODE_STROKE}" stroke-width="1.2"', labels["legend_seq"]))
    if any(n.kind == "ext" for n in nodes.values()):
        items.append(("swatch", 'fill="#ffffff" stroke="#A9B7BF" stroke-width="1.2" stroke-dasharray="4 3"', labels["legend_ext"]))
    if any(n.kind == "missing" for n in nodes.values()):
        items.append(("swatch", 'fill="#F7F9FA" stroke="#C2CCD1" stroke-width="1.2" stroke-dasharray="4 3"', labels["legend_missing"]))
    items.append(("arrow", "", labels["legend_calls"]))
    for kind, style, text in items:
        if kind == "swatch":
            parts.append(f'<rect x="{lx:.0f}" y="{ly:.0f}" width="22" height="14" rx="4" {style}/>')
            lx += 28
        else:
            parts.append(
                f'<path d="M {lx:.0f} {ly + 7:.0f} L {lx + 20:.0f} {ly + 7:.0f}" stroke="#{C_EDGE}" '
                'stroke-width="1.8" marker-end="url(#arrow)"/>'
            )
            lx += 26
        parts.append(f'<text x="{lx:.0f}" y="{ly + 11.5:.0f}" font-size="11.5" fill="#{C_MUTED}">{xml_escape(text)}</text>')
        lx += 7.2 * len(text) + 30
    parts.append("</svg>")
    return "".join(parts)


def render_svg_to_png(svg, W, H, png_path, scale=2, browser=None):
    exe = browser or next((p for p in BROWSER_CANDIDATES if os.path.isfile(p)), None)
    if not exe:
        raise RuntimeError("No Chromium browser found for diagram rendering. Checked: "
                           + "; ".join(BROWSER_CANDIDATES))
    tmp = tempfile.mkdtemp(prefix="tsdoc_")
    try:
        html_path = os.path.join(tmp, "diagram.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write("<!doctype html><html><head><meta charset='utf-8'>"
                    "<style>html,body{margin:0;padding:0;background:#fff}svg{display:block}</style>"
                    "</head><body>" + svg + "</body></html>")
        cmd = [
            exe, "--headless=new", "--disable-gpu", "--no-first-run", "--hide-scrollbars",
            "--user-data-dir=" + os.path.join(tmp, "profile"),
            "--screenshot=" + png_path,
            f"--window-size={int(W * scale)},{int(H * scale)}",
            "file:///" + html_path.replace("\\", "/"),
        ]
        subprocess.run(cmd, capture_output=True, timeout=120, check=False)
        if not os.path.isfile(png_path) or os.path.getsize(png_path) == 0:
            raise RuntimeError("Headless browser did not produce the diagram PNG: " + png_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# --------------------------------------------------------------------------
# DOCX helpers
# --------------------------------------------------------------------------

def _strip_theme(style):
    """Remove theme font/color bindings so explicit values actually win."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for att in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn("w:" + att), None)
    color = rpr.find(qn("w:color"))
    if color is not None:
        for att in ("themeColor", "themeTint", "themeShade"):
            color.attrib.pop(qn("w:" + att), None)


def tune_style(style, size, color, bold=None, before=None, after=None):
    f = style.font
    f.name = FONT
    f.size = Pt(size)
    f.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        f.bold = bold
    _strip_theme(style)
    pf = getattr(style, "paragraph_format", None)
    if pf is not None:
        if before is not None:
            pf.space_before = Pt(before)
        if after is not None:
            pf.space_after = Pt(after)


def fmt_run(run, size, color, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    return run


def add_rule(doc, color=C_RULE, sz=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)
    return p


def add_field(paragraph, instruction, placeholder=None, placeholder_fmt=None):
    """Insert a Word field (TOC, PAGE, ...) via fldChar runs."""
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    r._r.append(instr)
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)
    if placeholder:
        r = paragraph.add_run(placeholder)
        if placeholder_fmt:
            placeholder_fmt(r)
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def set_table_borders(table):
    """Horizontal lines only — no vertical grid — for a light, modern table."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag, val, sz, color in (
        ("top", "nil", "0", "auto"),
        ("left", "nil", "0", "auto"),
        ("right", "nil", "0", "auto"),
        ("bottom", "single", "4", C_TBL_LINE2),
        ("insideH", "single", "4", C_TBL_LINE),
        ("insideV", "nil", "0", "auto"),
    ):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_table_borders_minimal(table):
    """Only a light line under the whole table — for the compact step listing."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag, val, sz, color in (
        ("top", "nil", "0", "auto"),
        ("left", "nil", "0", "auto"),
        ("right", "nil", "0", "auto"),
        ("bottom", "single", "4", C_TBL_LINE),
        ("insideH", "nil", "0", "auto"),
        ("insideV", "nil", "0", "auto"),
    ):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_fixed_layout(table):
    """Enforce fixed column widths (Word honors the cell widths verbatim
    instead of auto-resizing columns to content)."""
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_cell_nowrap(cell):
    """Prevent Word from wrapping this cell's text onto a second line."""
    tcpr = cell._tc.get_or_add_tcPr()
    nw = OxmlElement("w:noWrap")
    tcpr.append(nw)


def set_cell_margins(table, top=50, bottom=50, left=110, right=110):
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def put_cell(cell, text, size=9.5, color=C_INK, bold=False, italic=False,
             indent=None, strike=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = indent
    r = fmt_run(p.add_run(text), size, color, bold=bold, italic=italic)
    if strike:
        r.font.strike = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------

def describe_calls(seq, file_base, labels):
    """One compact 'calls' line per sequence: internal, external and counts."""
    parts, unresolved = [], []
    merged = {}
    for c in seq.get("calls", []):
        if c.get("unresolved"):
            unresolved.append(c.get("step_name") or c.get("target_sequence") or "?")
            continue
        tseq = (c.get("target_sequence") or "").strip()
        if not tseq:
            continue
        tfile = (c.get("target_file") or "").strip()
        if tfile and os.path.basename(tfile).lower() == file_base:
            tfile = ""
        key = (os.path.basename(tfile), tseq)
        merged[key] = merged.get(key, 0) + max(1, int(c.get("count", 1) or 1))
    for (tfile, tseq), count in merged.items():
        label = tseq if not tfile else f"{tfile} → {tseq}"
        if not tfile and tseq == seq["name"]:
            label += " (" + labels["recursive"] + ")"
        if count > 1:
            label += f" (×{count})"
        parts.append(label)
    return parts, unresolved


def build_document(data, labels, png_path, diagram_w, diagram_h, out_path, tinter=None):
    doc = Document()

    # A4 with generous but businesslike margins.
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)

    tune_style(doc.styles["Normal"], 10.5, C_INK, after=6)
    tune_style(doc.styles["Heading 1"], 15, C_ACCENT, bold=True, before=18, after=8)
    tune_style(doc.styles["Heading 2"], 12.5, C_ACCENT2, bold=True, before=14, after=5)

    # --- Header block: title, meta line, rule, short description -----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    fmt_run(p.add_run(data["title"]), 26, C_ACCENT, bold=True)

    meta_bits = [data["file"]["path"]]
    if data["file"].get("version"):
        meta_bits.append(f"{labels['version']} {data['file']['version']}")
    meta_bits.append(f"{len(data['sequences'])} {labels['sequences_count']}")
    if data.get("generated"):
        meta_bits.append(data["generated"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    fmt_run(p.add_run("  ·  ".join(meta_bits)), 9, C_MUTED)
    add_rule(doc)

    file_base = os.path.basename(data["file"]["path"]).lower()

    desc = (data["file"].get("description") or "").strip()
    p = doc.add_paragraph()
    if desc:
        fmt_run(p.add_run(desc), 10.5, C_INK)
    else:
        fmt_run(p.add_run(labels["no_description"]), 10.5, C_MUTED, italic=True)

    # --- Sequence dependencies (overview, directly after the summary) --------
    # Placed up front as an at-a-glance architecture picture. Its heading is a
    # styled paragraph (not a Heading style) so it stays out of the TOC, which
    # then lists only the detailed per-sequence sections that follow.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    fmt_run(p.add_run(labels["dependencies"]), 15, C_ACCENT, bold=True)

    has_edges = any(s.get("calls") and any(not c.get("unresolved") for c in s["calls"])
                    for s in data["sequences"])
    ext = any(
        (c.get("target_file") or "").strip()
        and os.path.basename(c["target_file"]).lower() != file_base
        for s in data["sequences"] for c in s.get("calls", []) if not c.get("unresolved")
    )
    intro = labels["dep_intro"].format(ext=labels["dep_intro_ext"] if ext else "") \
        if has_edges else labels["dep_none"]
    fmt_run(doc.add_paragraph().add_run(intro), 10.5, C_INK)

    # Fit the diagram (rendered at 2x) into the portrait text area.
    usable_w = sec.page_width - sec.left_margin - sec.right_margin
    usable_h = sec.page_height - sec.top_margin - sec.bottom_margin - Cm(2.0)
    natural_w = Emu(int(diagram_w * 9525))   # 1 CSS px @96dpi = 9525 EMU
    natural_h = Emu(int(diagram_h * 9525))
    factor = min(1.0, usable_w / natural_w, usable_h / natural_h)
    doc.add_picture(png_path, width=Emu(int(natural_w * factor)))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Table of contents ---------------------------------------------------
    doc.add_page_break()
    p = doc.add_paragraph()  # deliberately NOT a Heading style: keeps it out of the TOC
    p.paragraph_format.space_after = Pt(10)
    fmt_run(p.add_run(labels["toc"]), 15, C_ACCENT, bold=True)
    add_field(
        doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
        placeholder=labels["toc_placeholder"],
        placeholder_fmt=lambda r: fmt_run(r, 9.5, C_MUTED, italic=True),
    )
    doc.add_page_break()

    # --- One section per sequence -------------------------------------------
    doc.add_heading(labels["sequences"], level=1)
    col_widths = (3.4, 2.4, 3.0, 2.6, 4.8)
    headers = (labels["param_name"], labels["param_type"], labels["param_passing"],
               labels["param_default"], labels["param_desc"])

    for seq in data["sequences"]:
        doc.add_heading(seq["name"], level=2)
        sdesc = (seq.get("description") or "").strip()
        p = doc.add_paragraph()
        if sdesc:
            fmt_run(p.add_run(sdesc), 10.5, C_INK)
        else:
            fmt_run(p.add_run(labels["no_description"]), 10.5, C_MUTED, italic=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        fmt_run(p.add_run(labels["parameters"]), 10, C_ACCENT, bold=True)

        params = seq.get("parameters", [])
        if not params:
            p = doc.add_paragraph()
            fmt_run(p.add_run(labels["no_parameters"]), 9.5, C_MUTED, italic=True)
        else:
            table = doc.add_table(rows=1, cols=5)
            table.autofit = False
            set_table_borders(table)
            set_cell_margins(table)
            hdr = table.rows[0]
            mark_header_row(hdr)
            hdr.height = Cm(0.68)
            hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell, text, w in zip(hdr.cells, headers, col_widths):
                cell.width = Cm(w)
                set_cell_shading(cell, C_ACCENT)
                put_cell(cell, text, size=9.5, color="FFFFFF", bold=True)
            for prm in params:
                by_ref = str(prm.get("by_ref", "")).strip().lower() in ("true", "1", "yes", "byref", "by reference")
                row = table.add_row()
                values = (
                    prm.get("name", ""),
                    prm.get("type", ""),
                    labels["by_ref"] if by_ref else labels["by_val"],
                    str(prm.get("default", "")) if str(prm.get("default", "")).strip() else "—",
                    (prm.get("description") or "").strip() or "—",
                )
                for cell, text, w in zip(row.cells, values, col_widths):
                    cell.width = Cm(w)
                    put_cell(cell, text, size=9.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)

        # Compact step listing, grouped Setup / Main / Cleanup, flow-indented.
        # No step-type column (user preference) — just the (icon +) step name
        # and, if present, a detail column. Column widths are derived from the
        # actual content so lines don't wrap when they can fit; only content
        # that genuinely exceeds the page width wraps.
        groups = seq.get("steps") or {}
        ordered = [(g, groups.get(g) or []) for g in ("Setup", "Main", "Cleanup")]
        ordered = [(g, s) for g, s in ordered if s]
        if ordered:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            fmt_run(p.add_run(labels["steps"]), 10, C_ACCENT, bold=True)

            has_detail = any((st.get("detail") or "").strip()
                             for _, steps in ordered for st in steps)

            # Estimate the width each column needs to avoid wrapping (cm).
            USABLE = 16.6            # A4 text width: 21 - 2*2.2 cm
            ICON_GAP = 0.52 if (tinter and tinter.ok) else 0.0
            PAD = 0.34              # cell left+right margins
            name_need = detail_need = 0.0
            for _, steps in ordered:
                for depth, st in step_rows(steps):
                    t = (st.get("type") or "").strip()
                    strong = t in FLOW_ALL and t not in FLOW_END
                    chw = 0.170 if strong else 0.158
                    name_need = max(name_need, 0.32 * depth + ICON_GAP
                                    + len(st.get("name", "")) * chw + PAD)
                    d = (st.get("detail") or "").strip()
                    if d:
                        detail_need = max(detail_need, len(d) * 0.150 + PAD)

            if has_detail:
                if name_need + detail_need <= USABLE:
                    widths, fits = (name_need, detail_need), True   # natural, no wrap
                else:
                    name_w = min(name_need, USABLE * 0.60)
                    widths, fits = (name_w, USABLE - name_w), False  # detail may wrap
            else:
                widths, fits = (min(name_need, USABLE),), name_need <= USABLE

            table = doc.add_table(rows=0, cols=len(widths))
            table.autofit = False
            set_fixed_layout(table)
            set_table_borders_minimal(table)
            set_cell_margins(table, top=4, bottom=4, left=96, right=96)

            def compact(paragraph, indent=None):
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(12)
                if indent:
                    pf.left_indent = indent

            for group, steps in ordered:
                row = table.add_row()
                merged = row.cells[0].merge(row.cells[-1])
                set_cell_shading(merged, "E8EFF2")
                compact(merged.paragraphs[0])
                fmt_run(merged.paragraphs[0].add_run(group), 9, C_ACCENT, bold=True)
                merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for depth, st in step_rows(steps):
                    t = (st.get("type") or "").strip()
                    if t in FLOW_END:
                        color, bold = C_MUTED, False
                    elif t in FLOW_ALL:
                        color, bold = C_ACCENT2, True
                    else:
                        color, bold = C_INK, False
                    disabled = st.get("enabled") is False
                    cells = table.add_row().cells
                    p = cells[0].paragraphs[0]
                    compact(p, indent=Cm(0.32 * depth) if depth else None)
                    icon_color = C_MUTED if (disabled or t in FLOW_END) else C_ACCENT
                    icon = tinter.png(t, icon_color) if tinter else None
                    if icon:
                        p.add_run().add_picture(icon, height=Pt(9.5))
                    r = fmt_run(p.add_run(("  " if icon else "") + st.get("name", "")),
                                9, C_MUTED if disabled else color,
                                bold=bold and not disabled)
                    if disabled:
                        r.font.strike = True
                    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if fits:
                        set_cell_nowrap(cells[0])
                    if has_detail:
                        compact(cells[1].paragraphs[0])
                        fmt_run(cells[1].paragraphs[0].add_run((st.get("detail") or "").strip()),
                                8.5, C_MUTED)
                        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        if fits:
                            set_cell_nowrap(cells[1])
                    for cell, w in zip(cells, widths):
                        cell.width = Cm(w)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)

        calls, unresolved = describe_calls(seq, file_base, labels)
        if calls:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            fmt_run(p.add_run(labels["calls"] + ": "), 9.5, C_MUTED, bold=True)
            fmt_run(p.add_run(", ".join(calls)), 9.5, C_MUTED)
        if unresolved:
            p = doc.add_paragraph()
            fmt_run(p.add_run(labels["unresolved_calls"] + ": "), 9.5, C_MUTED, bold=True)
            fmt_run(p.add_run(", ".join(unresolved)), 9.5, C_MUTED, italic=True)

    # --- Footer: right-aligned page numbers -----------------------------------
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fmt_run(footer_p.add_run(labels["page"] + " "), 8.5, C_MUTED)
    add_field(footer_p, "PAGE")
    fmt_run(footer_p.add_run(" " + labels["page_of"] + " "), 8.5, C_MUTED)
    add_field(footer_p, "NUMPAGES")
    for run in footer_p.runs:
        fmt_run(run, 8.5, C_MUTED, bold=run.font.bold or False)

    # Ask Word to refresh fields (TOC, page counts) when the document opens.
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)

    doc.save(out_path)


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def main(argv):
    args = [a for a in argv if not a.startswith("--")]
    if len(args) < 2:
        print(__doc__)
        return 1
    data_path, out_path = os.path.abspath(args[0]), os.path.abspath(args[1])

    def opt(name):
        return argv[argv.index(name) + 1] if name in argv else None

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    labels = dict(LABELS.get(str(data.get("language", "de")).lower(), LABELS["de"]))
    labels.update(data.get("labels", {}))

    nodes, order, edges, back = build_graph(data, labels)
    W, H, margin, route_pad = layout_graph(nodes, order, edges, back)
    svg = make_svg(nodes, order, edges, back, W, H, margin, labels, route_pad=route_pad)

    png_path = os.path.abspath(opt("--diagram-out") or os.path.join(
        tempfile.gettempdir(), "tsdoc_diagram.png"))
    render_svg_to_png(svg, W, H, png_path, browser=opt("--browser"))

    components = find_teststand_components(opt("--teststand-dir"))
    icon_dir = tempfile.mkdtemp(prefix="tsdoc_icons_") if (components and Image) else None
    tinter = IconTinter(components, icon_dir)
    try:
        build_document(data, labels, png_path, W, H, out_path, tinter)
    finally:
        if icon_dir:
            shutil.rmtree(icon_dir, ignore_errors=True)

    n_params = sum(len(s.get("parameters", [])) for s in data["sequences"])
    n_unres = sum(1 for s in data["sequences"] for c in s.get("calls", []) if c.get("unresolved"))
    print(f"[ok] docx    : {out_path}")
    print(f"[ok] diagram : {png_path} ({int(W * 2)}x{int(H * 2)} px, logical {int(W)}x{int(H)})")
    if tinter.ok:
        print(f"[ok] icons   : original TestStand step icons, tinted ({components})")
    else:
        print("[--] icons   : omitted (TestStand installation or Pillow not found)")
    print(f"[ok] content : {len(data['sequences'])} sequences, {n_params} parameters, "
          f"{len(edges)} call edges ({len(back)} recursive), {n_unres} unresolved calls")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
